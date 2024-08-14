import os
import slicer
from slicer.ScriptedLoadableModule import *
import logging
import ctk  # Import ctk for UI components
import qt  # Import qt for Qt-based UI components
import pandas as pd  # Import pandas for DataFrame operations
import uuid  # Import uuid for generating unique IDs
import shutil  # Import shutil to check for pandoc installation

class Medical_Data_Anonymizer_Module(ScriptedLoadableModule):
    def __init__(self, parent):
        ScriptedLoadableModule.__init__(self, parent)
        self.parent.title = "Medical Data Anonymizer Module"
        self.parent.categories = ["Examples"]
        self.parent.dependencies = []
        self.parent.contributors = ["Jonas Bianchi"]
        self.parent.helpText = """This module anonymizes text files."""
        self.parent.acknowledgementText = """Developed using Slicer resources."""

class Medical_Data_Anonymizer_ModuleWidget(ScriptedLoadableModuleWidget):

    def setup(self):
        ScriptedLoadableModuleWidget.setup(self)

        # Add labels for input and output directories
        self.inputLabel = qt.QLabel("Files to be Anonymized")
        self.layout.addWidget(self.inputLabel)

        self.inputDirectoryButton = ctk.ctkPathLineEdit()
        self.inputDirectoryButton.filters = ctk.ctkPathLineEdit.Dirs
        self.inputDirectoryButton.setToolTip("Select the directory containing the files to be anonymized.")
        self.layout.addWidget(self.inputDirectoryButton)

        self.outputLabel = qt.QLabel("Output Anonymized Files")
        self.layout.addWidget(self.outputLabel)

        self.outputDirectoryButton = ctk.ctkPathLineEdit()
        self.outputDirectoryButton.filters = ctk.ctkPathLineEdit.Dirs
        self.outputDirectoryButton.setToolTip("Select the directory to save the anonymized files.")
        self.layout.addWidget(self.outputDirectoryButton)

        # Button to install dependencies
        self.installDependenciesButton = qt.QPushButton("Install Dependencies")
        self.installDependenciesButton.toolTip = "Install required dependencies."
        self.layout.addWidget(self.installDependenciesButton)
        self.installDependenciesButton.connect('clicked(bool)', self.install_dependencies)

        self.anonymizeButton = qt.QPushButton("Anonymize")
        self.anonymizeButton.toolTip = "Run the anonymization."
        self.layout.addWidget(self.anonymizeButton)
        self.anonymizeButton.connect('clicked(bool)', self.onAnonymizeButton)

        # Add vertical spacer
        self.layout.addStretch(1)

    def install_dependencies(self):
        slicer.util.pip_install('spacy')
        slicer.util.pip_install('pandas')
        slicer.util.pip_install('python-docx')  # Install python-docx for handling .docx files

        # Download spaCy language model if not already downloaded
        import spacy
        spacy.cli.download("en_core_web_sm")

        # Notify user to restart Slicer
        qt.QMessageBox.warning(
            slicer.util.mainWindow(),
            'Restart Required',
            'Dependencies have been installed. Please restart 3D Slicer to complete the installation.'
        )

    def onAnonymizeButton(self):
        try:
            import spacy
            import docx  # Ensure that docx is imported before using it
        except ImportError:
            qt.QMessageBox.warning(
                slicer.util.mainWindow(),
                'Dependencies Not Installed',
                'Please install the dependencies first by clicking the "Install Dependencies" button and restart 3D Slicer before running the anonymization.'
            )
            return

        # Continue with the anonymization process if dependencies are installed
        nlp = spacy.load("en_core_web_sm")
        input_folder = self.inputDirectoryButton.currentPath
        output_folder = self.outputDirectoryButton.currentPath
        csv_file_path = os.path.join(output_folder, "file_mappings.csv")

        # Run the anonymization process
        birth_date_keywords = ["born", "date of birth", "dob", "b.o.b"]
        file_mappings = []

        for root, dirs, files in os.walk(input_folder):
            for file in files:
                if file.startswith("~$"):  # Skip temporary or hidden files
                    continue

                if file.endswith(".docx"):
                    input_file_path = os.path.join(root, file)

                    try:
                        unique_id = str(uuid.uuid4())
                        doc = docx.Document(input_file_path)  # Ensure we're using the imported docx module
                        full_text = "\n".join([para.text for para in doc.paragraphs])
                        anonymized_text = self.anonymize_text(full_text, nlp, birth_date_keywords)

                        new_doc = docx.Document()
                        for line in anonymized_text.split("\n"):
                            new_doc.add_paragraph(line)

                        new_file_name = f"{unique_id}.docx"
                        output_docx_path = os.path.join(output_folder, new_file_name)
                        new_doc.save(output_docx_path)

                        file_mappings.append({
                            "Original File Name": file,
                            "Anonymized File Name (DOCX)": new_file_name,
                        })

                        logging.info(f"Anonymized file created: {output_docx_path}")

                    except Exception as e:
                        logging.error(f"Error processing {file}: {e}")

        mappings_df = pd.DataFrame(file_mappings)
        if not mappings_df.empty:
            mappings_df.to_csv(csv_file_path, index=False)
            logging.info(f"Anonymization complete. File mappings saved to {csv_file_path}.")
        else:
            logging.info("No valid files were processed. CSV file not created.")

    def anonymize_text(self, text, nlp, birth_date_keywords):
        doc = nlp(text)
        anonymized_text = text
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                anonymized_text = anonymized_text.replace(ent.text, "ANONYMOUS")
            elif ent.label_ == "DATE":
                context = text[max(0, ent.start_char-50):ent.end_char+50].lower()
                if any(keyword in context for keyword in birth_date_keywords):
                    anonymized_text = anonymized_text.replace(ent.text, "ANONYMOUS DATE OF BIRTH")
        return anonymized_text
